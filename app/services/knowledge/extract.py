"""
Bóc text thô từ file admin upload (PDF / DOCX / XLSX / MD) → list đoạn text để chunk.

Mỗi định dạng có cách bóc khác nhau:
  • PDF  : PyMuPDF (fitz) — đọc text từng trang (tái dùng lib đã có ở homework/extract).
  • DOCX : python-docx — đọc từng paragraph.
  • XLSX : openpyxl — mỗi HÀNG = 1 đơn vị (bảng Q&A: ghép các ô thành 1 đoạn).
  • MD   : text thuần — tách khối theo dòng trống, giống hệt update_document_content
    (sửa tay trong CMS) vì bản chất là cùng 1 loại nội dung.

Trả về list[str] các khối text THÔ (chưa chunk). Bước chunk (chunk.py) lo cắt tiếp
khối văn xuôi dài; riêng XLSX mỗi hàng đã là 1 đơn vị tự nhiên nên trả sẵn từng hàng.
"""
from __future__ import annotations

import io
import re


class UnsupportedFileType(Exception):
    pass


def detect_source_type(file_name: str) -> str:
    name = file_name.lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".xlsx") or name.endswith(".xls"):
        return "xlsx"
    if name.endswith(".md") or name.endswith(".markdown"):
        return "md"
    raise UnsupportedFileType(f"Định dạng không hỗ trợ: {file_name} (chỉ pdf/docx/xlsx/md)")


def extract_blocks(file_bytes: bytes, source_type: str) -> list[str]:
    """Bóc text → list khối. Rỗng nếu file không có text đọc được."""
    if source_type == "pdf":
        return _extract_pdf(file_bytes)
    if source_type == "docx":
        return _extract_docx(file_bytes)
    if source_type == "xlsx":
        return _extract_xlsx(file_bytes)
    if source_type == "md":
        return _extract_md(file_bytes)
    raise UnsupportedFileType(source_type)


def _extract_md(file_bytes: bytes) -> list[str]:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    return [b for b in re.split(r"\n\s*\n", text) if b.strip()]


def _extract_pdf(file_bytes: bytes) -> list[str]:
    import fitz  # PyMuPDF (đã có trong requirements)

    blocks: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                blocks.append(text)
    return blocks


def _extract_docx(file_bytes: bytes) -> list[str]:
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(file_bytes))
    # Gộp paragraph liền nhau thành khối; đoạn rỗng = ranh giới khối (giữ cấu trúc tài liệu).
    blocks: list[str] = []
    buf: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            buf.append(t)
        elif buf:
            blocks.append("\n".join(buf))
            buf = []
    if buf:
        blocks.append("\n".join(buf))

    # Bảng trong docx: mỗi hàng ghép ô thành 1 dòng (thường là bảng Q&A/thông số).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" — ".join(cells))
    return blocks


def _extract_xlsx(file_bytes: bytes) -> list[str]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    blocks: list[str] = []
    for ws in wb.worksheets:
        rows = ws.iter_rows(values_only=True)
        header = next(rows, None)
        # Nếu có header (vd "Câu hỏi | Trả lời"), ghép "cột: giá trị" cho mỗi hàng để
        # đoạn tự đủ nghĩa; không có header thì nối các ô bằng " — ".
        header_labels = [str(h).strip() if h is not None else "" for h in header] if header else []
        has_header = any(header_labels)
        for row in rows:
            values = [("" if v is None else str(v).strip()) for v in row]
            if not any(values):
                continue
            if has_header:
                parts = [
                    f"{header_labels[i]}: {values[i]}"
                    for i in range(min(len(header_labels), len(values)))
                    if values[i]
                ]
                blocks.append(". ".join(parts))
            else:
                blocks.append(" — ".join(v for v in values if v))
    wb.close()
    return blocks
